import streamlit as st
import os
import random
from main import generate_quiz
from fpdf import FPDF
from docx import Document
import io

# -------------------- PAGE CONFIG --------------------
st.set_page_config(page_title="🧠 Cognitive Quiz Generator", layout="centered")

# -------------------- CSS STYLING --------------------
st.markdown("""
<style>
.stApp {
    background: linear-gradient(135deg, #f6d365, #fda085, #a1c4fd, #c2e9fb);
    background-attachment: fixed;
}
body {
    font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
}
h1, h2, h3 { color: #1f2937; }
.quiz-card {
    background: linear-gradient(135deg, #ffecd2, #fcb69f);
    border-radius: 15px;
    padding: 18px;
    margin: 10px 0;
    box-shadow: 0px 6px 12px rgba(0,0,0,0.15);
    transition: transform 0.3s ease-in-out, box-shadow 0.3s ease-in-out;
}
.quiz-card:hover {
    transform: translateY(-5px);
    box-shadow: 0px 12px 20px rgba(0,0,0,0.25);
}
.stButton>button {
    background: linear-gradient(90deg, #ff7e5f, #feb47b);
    color: white;
    border-radius: 12px;
    padding: 10px 25px;
    font-weight: bold;
    font-size: 16px;
    transition: all 0.3s ease-in-out;
}
.stButton>button:hover {
    background: linear-gradient(90deg, #feb47b, #ff7e5f);
    transform: scale(1.05);
}
.stNumberInput>div>div>input {
    border-radius: 10px;
    padding: 7px;
}
.stRadio>div>div>label {
    padding: 6px 12px;
    background-color: #ffffffb3;
    border-radius: 10px;
    margin: 2px;
    transition: all 0.2s ease-in-out;
}
.stRadio>div>div>label:hover {
    background-color: #ffffff;
}
</style>
""", unsafe_allow_html=True)

# -------------------- SIDEBAR INSTRUCTIONS --------------------
st.sidebar.markdown("""
<div style="background-color: #fff3cd; border-left: 6px solid #ffeeba; padding: 15px; border-radius: 10px;">
<h3 style="color: #856404;">📌 Instructions</h3>
<ul style="color: #856404; font-size: 14px;">
<li>Upload a text, Word (.docx), or PDF file containing the content you want to quiz on.</li>
<li>Select the number of questions for the quiz.</li>
<li>Choose the quiz format: Cloze (fill in the blanks) or MCQ (multiple choice).</li>
<li>Click <strong>Generate Quiz</strong> to create a new set of questions.</li>
<li>Answer the questions and click <strong>Submit Quiz</strong> to see your score and review answers.</li>
<li>Each time you generate, a new set of questions will appear for better learning.</li>
<li>You can download the generated quiz as PDF or Word for offline study.</li>
</ul>
</div>
""", unsafe_allow_html=True)

# -------------------- PAGE TITLE --------------------
st.title("🧠 Cognitive Quiz Generator")

# -------------------- INPUTS --------------------
uploaded_file = st.file_uploader("Upload a text, Word, or PDF file", type=["txt", "docx", "pdf"])
num_questions = st.number_input("Number of questions", min_value=1, max_value=20, value=5)
quiz_type = st.radio("Choose quiz format", ["Cloze", "MCQ"])

# -------------------- SESSION STATE --------------------
if "quiz" not in st.session_state:
    st.session_state.quiz = None
if "answers" not in st.session_state:
    st.session_state.answers = {}

# -------------------- GENERATE QUIZ --------------------
if uploaded_file and st.button("Generate Quiz"):
    temp_path = "temp_file" + os.path.splitext(uploaded_file.name)[1].lower()
    with open(temp_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    with st.spinner("Generating quiz..."):
        quiz = generate_quiz(temp_path, num_questions=num_questions)

    # Shuffle MCQ options
    for q in quiz.get("mcq", []):
        opts = q.get("options")
        if not opts or not isinstance(opts, list):
            opts = [q.get("answer", "")]
        random.shuffle(opts)
        q["options"] = opts

    st.session_state.quiz = quiz
    st.session_state.answers = {}

    # Display generation summary
    cloze_count = len(quiz.get("cloze", []))
    mcq_count = len(quiz.get("mcq", []))

    st.write("### Quiz Generation Results")
    col1, col2 = st.columns(2)

    with col1:
        st.metric("Requested Questions", num_questions)
    with col2:
        if quiz_type == "Cloze":
            st.metric("Generated Cloze", cloze_count)
        else:
            st.metric("Generated MCQ", mcq_count)

    if quiz_type == "Cloze":
        if cloze_count == num_questions:
            st.success(f"✅ Perfect! Generated all {cloze_count} cloze questions.")
        elif cloze_count > 0:
            st.warning(f"⚠️ Generated {cloze_count} cloze questions out of {num_questions} requested.")
        else:
            st.error("❌ Could not generate any cloze questions.")
    else:
        if mcq_count == num_questions:
            st.success(f"✅ Perfect! Generated all {mcq_count} MCQ questions.")
        elif mcq_count > 0:
            st.warning(f"⚠️ Generated {mcq_count} MCQ questions out of {num_questions} requested.")
        else:
            st.error("❌ Could not generate any MCQ questions.")

    if cloze_count > 0 or mcq_count > 0:
        st.info("📝 Scroll down to attempt the quiz!")

    # Clean up temp file
    try:
        os.remove(temp_path)
    except:
        pass

# -------------------- RENDER QUIZ --------------------
if st.session_state.quiz:
    quiz = st.session_state.quiz

    # Cloze
    if quiz_type == "Cloze":
        st.subheader("✍️ Cloze Quiz")
        for i, q in enumerate(quiz.get("cloze", [])[:num_questions], start=1):
            st.markdown(f'<div class="quiz-card"><strong>Q{i}.</strong> {q["question"]}</div>', unsafe_allow_html=True)
            user_in = st.text_input(f"Your answer for Q{i}", key=f"cloze_{i}")
            st.session_state.answers[f"cloze_{i}"] = {"user": user_in.strip(), "correct": q["answer"]}

    # MCQ
    else:
        st.subheader("🎯 MCQ Quiz")
        for i, q in enumerate(quiz.get("mcq", [])[:num_questions], start=1):
            st.markdown(f'<div class="quiz-card"><strong>Q{i}.</strong> {q["question"]}</div>', unsafe_allow_html=True)
            opts = q.get("options") or [q.get("answer", "")]
            placeholder = "-- Select an answer --"
            opts_with_placeholder = [placeholder] + opts
            user_sel = st.selectbox(f"Choose answer for Q{i}", options=opts_with_placeholder, index=0, key=f"mcq_{i}")
            st.session_state.answers[f"mcq_{i}"] = {"user": "" if user_sel == placeholder else user_sel, "correct": q["answer"]}

# -------------------- SUBMIT QUIZ --------------------
if st.session_state.quiz and st.button("Submit Quiz"):
    answers = st.session_state.answers
    correct = 0
    total = 0

    relevant = {k: v for k, v in answers.items() if (k.startswith("cloze_") if quiz_type == "Cloze" else k.startswith("mcq_"))}
    for v in relevant.values():
        total += 1
        if v["user"].strip().lower() == v["correct"].strip().lower():
            correct += 1

    percentage = (correct / total * 100) if total > 0 else 0
    st.success(f"✅ You scored {correct} out of {total} ({percentage:.1f}%)")

    st.write("### Review")
    for idx, v in enumerate(relevant.values(), start=1):
        user_ans = v["user"] or "(no answer)"
        is_correct = v["user"].strip().lower() == v["correct"].strip().lower()
        status_icon = "✅" if is_correct else "❌"
        st.markdown(f'''<div class="quiz-card">
        {status_icon} Q{idx}: Your answer = <strong>{user_ans}</strong> | Correct = <strong>{v["correct"]}</strong>
        </div>''', unsafe_allow_html=True)

# -------------------- EXPORT QUIZ --------------------
def sanitize_text(text):
    replacements = {
        "’": "'", "‘": "'", "“": '"', "”": '"',
        "—": "-", "…": "...", "🧠": "", "✍️": "", "🎯": ""
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    return text.encode("latin-1", "replace").decode("latin-1")

def create_pdf(quiz_data, num_questions):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, "Cognitive Quiz Generator", ln=True, align="C")
    pdf.ln(10)
    pdf.set_font("Arial", "", 12)

    cloze_questions = quiz_data.get("cloze", [])[:num_questions]
    mcq_questions = quiz_data.get("mcq", [])[:num_questions]

    if cloze_questions:
        pdf.cell(0, 8, f"Cloze Questions ({len(cloze_questions)}):", ln=True)
        pdf.ln(2)
        for i, q in enumerate(cloze_questions, start=1):
            pdf.multi_cell(0, 8, f"Q{i}. {sanitize_text(q['question'])}")
            pdf.ln(1)

    if mcq_questions:
        pdf.ln(5)
        pdf.cell(0, 8, f"MCQ Questions ({len(mcq_questions)}):", ln=True)
        pdf.ln(2)
        for i, q in enumerate(mcq_questions, start=1):
            pdf.multi_cell(0, 8, f"Q{i}. {sanitize_text(q['question'])}")
            pdf.ln(1)
            for idx, opt in enumerate(q.get("options", []), start=1):
                pdf.multi_cell(0, 8, f"   {chr(64+idx)}. {sanitize_text(opt)}")
            pdf.ln(2)

    return pdf.output(dest="S").encode("latin-1", "replace")

def create_word(quiz_data, num_questions):
    doc = Document()
    doc.add_heading("Cognitive Quiz Generator", 0)

    cloze_questions = quiz_data.get("cloze", [])[:num_questions]
    mcq_questions = quiz_data.get("mcq", [])[:num_questions]

    if cloze_questions:
        doc.add_heading(f"Cloze Questions ({len(cloze_questions)}):", level=1)
        for i, q in enumerate(cloze_questions, start=1):
            doc.add_paragraph(sanitize_text(f"Q{i}. {q['question']}"))

    if mcq_questions:
        doc.add_heading(f"MCQ Questions ({len(mcq_questions)}):", level=1)
        for i, q in enumerate(mcq_questions, start=1):
            doc.add_paragraph(sanitize_text(f"Q{i}. {q['question']}"))
            for idx, opt in enumerate(q.get("options", []), start=1):
                doc.add_paragraph(sanitize_text(f"   {chr(64+idx)}. {opt}"), style="List Bullet")

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

if st.session_state.quiz:
    st.markdown("---")
    st.subheader("📥 Download Quiz")

    cloze_count = len(st.session_state.quiz.get("cloze", []))
    mcq_count = len(st.session_state.quiz.get("mcq", []))

    pdf_bytes = create_pdf(st.session_state.quiz, num_questions)
    st.download_button("📄 Download as PDF", data=pdf_bytes, file_name="quiz.pdf", mime="application/pdf")
    word_file = create_word(st.session_state.quiz, num_questions)
    st.download_button("📝 Download as Word", data=word_file, file_name="quiz.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    st.info(f"Export includes: {cloze_count} cloze, {mcq_count} MCQ")
